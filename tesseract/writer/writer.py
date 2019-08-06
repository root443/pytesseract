import re

from tesseract.settings import *


def writing(filename, text):
    if re.search(r"\.docx", filename):
        writingdocx(filename=filename, content=text)
    else:
        writingtxt(filename=filename, content=text)

def writingdocx(filename, content, column=0):
    from docx import Document
    from docx.oxml.ns import qn

    document = Document()

    # rtlstyle = document.styles.add_style('rtl', WD_STYLE_TYPE.CHARACTER)
    paragraph = document.add_paragraph(text=content)
    # run = paragraph.add_run(text=content)
    # run.style = rtlstyle
    # run.font.rtl = True
    # paragraph = document.add_paragraph(content)
    if column > 0:
        section = document.sections[0]
        cols = section._sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), column)
    document.save('{d}{f}'.format(d=FOLDER_OUT, f=filename))

def writingtxt(filename, content):
    with open('{d}{f}.txt'.format(d=FOLDER_OUT, f=filename), 'wb') as f:
        f.write(content.encode())
